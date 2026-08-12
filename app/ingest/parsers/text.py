"""TXT / PDF / DOCX 解析（FR-T-001），包含表格和扫描页 OCR。"""

from __future__ import annotations

import re
from datetime import datetime
from decimal import Decimal
from pathlib import Path

from app.core.timeutil import BUSINESS_TZ, ensure_aware
from app.ingest.parsers.base import ParsedDocument, ParseError, RawSegment

# 样例投研资料的标题行：[DOC-DEMO-001 | 内部研究摘要 | 2026-01-15 09:00]
_HEADER = re.compile(r"^\[(?P<doc_id>[^|\]]+)\|(?P<doc_type>[^|\]]+)\|(?P<published>[^|\]]+)\]$")


def _parse_published(text: str) -> datetime | None:
    text = text.strip()
    for fmt in ("%Y-%m-%d %H:%M", "%Y-%m-%d"):
        try:
            return ensure_aware(datetime.strptime(text, fmt), assume=BUSINESS_TZ)
        except ValueError:
            continue
    return None


def parse_txt(text: str) -> ParsedDocument:
    """按空行分段。段落是引用定位的最小单位。"""
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if not paragraphs:
        raise ParseError("文件内容为空，无法解析")

    return ParsedDocument(
        title=paragraphs[0][:200],
        segments=[RawSegment(ordinal=i, content=p) for i, p in enumerate(paragraphs, start=1)],
    )


def parse_sample_pack(text: str) -> list[tuple[str, ParsedDocument]]:
    """解析样例投研资料，一个文件含多份文档。

    返回 (document_id, ParsedDocument)。published_at 取标题行时间——这是收益标签
    的时间起点（FLD-002），不能用入库时间代替。
    """
    results: list[tuple[str, ParsedDocument]] = []
    doc_id: str | None = None
    doc_type: str | None = None
    published: datetime | None = None
    buffer: list[str] = []

    def flush() -> None:
        if doc_id is None or not buffer:
            return
        results.append(
            (
                doc_id,
                ParsedDocument(
                    title=f"{doc_type} {doc_id}" if doc_type else doc_id,
                    doc_type=doc_type,
                    published_at=published,
                    segments=[
                        RawSegment(ordinal=i, content=p) for i, p in enumerate(buffer, start=1)
                    ],
                ),
            )
        )

    for raw in text.splitlines():
        line = raw.strip()
        matched = _HEADER.match(line)
        if matched:
            flush()
            doc_id = matched.group("doc_id").strip()
            doc_type = matched.group("doc_type").strip()
            published = _parse_published(matched.group("published"))
            buffer = []
            continue
        if doc_id is not None and line:
            buffer.append(line)

    flush()
    return results


def parse_pdf(path: Path) -> ParsedDocument:
    """逐页解析文本和表格；无文本的扫描页转图片后走本地 OCR。"""
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - 依赖缺失属环境问题
        raise ParseError(f"PDF 解析依赖缺失: {exc}", recoverable=False) from exc

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise ParseError(f"PDF 无法打开: {exc}") from exc

    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover
        raise ParseError(f"PDF 表格解析依赖缺失: {exc}", recoverable=False) from exc

    segments: list[RawSegment] = []
    ordinal = 0
    with pdfplumber.open(path) as plumber:
        for page_no, page in enumerate(reader.pages, start=1):
            native_text = page.extract_text() or ""
            for para in re.split(r"\n\s*\n", native_text):
                cleaned = para.strip()
                if cleaned:
                    ordinal += 1
                    segments.append(RawSegment(ordinal=ordinal, content=cleaned, page=page_no))

            for table_no, table in enumerate(plumber.pages[page_no - 1].extract_tables(), start=1):
                for row_no, row in enumerate(table or [], start=1):
                    cells = [re.sub(r"\s+", " ", cell or "").strip() for cell in row]
                    if not any(cells):
                        continue
                    ordinal += 1
                    end_column = _excel_column(len(cells))
                    segments.append(
                        RawSegment(
                            ordinal=ordinal,
                            content=" | ".join(cells),
                            page=page_no,
                            content_kind="table_row",
                            extraction_method="native",
                            table_index=table_no,
                            row_index=row_no,
                            cell_range=f"A{row_no}:{end_column}{row_no}",
                        )
                    )

            if not native_text.strip():
                for text, confidence in _ocr_pdf_page(path, page_no - 1):
                    ordinal += 1
                    segments.append(
                        RawSegment(
                            ordinal=ordinal,
                            content=text,
                            page=page_no,
                            extraction_method="ocr",
                            confidence=confidence,
                        )
                    )

    if not segments:
        raise ParseError("PDF 原生解析和 OCR 均未提取到可用文本")

    return ParsedDocument(title=path.stem, segments=segments)


def parse_docx(path: Path) -> ParsedDocument:
    """DOCX 解析。表格按行拼成段落，避免表内数字丢失上下文。"""
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise ParseError(f"DOCX 解析依赖缺失: {exc}", recoverable=False) from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise ParseError(f"DOCX 无法打开: {exc}") from exc

    segments: list[RawSegment] = []
    for para in document.paragraphs:
        if para.text.strip():
            segments.append(RawSegment(ordinal=len(segments) + 1, content=para.text.strip()))
    for table_no, table in enumerate(document.tables, start=1):
        for row_no, row in enumerate(table.rows, start=1):
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                segments.append(
                    RawSegment(
                        ordinal=len(segments) + 1,
                        content=" | ".join(cells),
                        content_kind="table_row",
                        table_index=table_no,
                        row_index=row_no,
                        cell_range=f"A{row_no}:{_excel_column(len(cells))}{row_no}",
                    )
                )

    if not segments:
        raise ParseError("DOCX 未提取到文本")

    return ParsedDocument(title=path.stem, segments=segments)


def parse_file(path: Path) -> ParsedDocument:
    """按扩展名分派。不支持的格式按失败处理，保留文件。"""
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return parse_txt(path.read_text(encoding="utf-8"))
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".docx":
        return parse_docx(path)
    raise ParseError(f"MVP 仅支持 PDF / DOCX / TXT，收到 {suffix or '无扩展名'}", recoverable=False)


def _excel_column(index: int) -> str:
    value = max(index, 1)
    result = ""
    while value:
        value, remainder = divmod(value - 1, 26)
        result = chr(65 + remainder) + result
    return result


def _ocr_pdf_page(path: Path, page_index: int) -> list[tuple[str, Decimal | None]]:
    """把单页渲染为图片并用 RapidOCR 提取按阅读顺序排列的文本行。"""
    try:
        import pypdfium2 as pdfium  # type: ignore[import-untyped]
        from rapidocr import RapidOCR
    except ImportError as exc:  # pragma: no cover
        raise ParseError(f"OCR 依赖缺失: {exc}", recoverable=False) from exc

    try:
        pdf = pdfium.PdfDocument(path)
        page = pdf[page_index]
        bitmap = page.render(scale=2.5)
        image = bitmap.to_pil()
        output: object = RapidOCR()(image)
        texts = list(getattr(output, "txts", None) or [])
        scores = list(getattr(output, "scores", None) or [])
        image.close()
        bitmap.close()
        page.close()
        pdf.close()
    except Exception as exc:
        raise ParseError(f"PDF 第 {page_index + 1} 页 OCR 失败: {exc}") from exc
    return [
        (text.strip(), Decimal(str(scores[index])) if index < len(scores) else None)
        for index, text in enumerate(texts)
        if text and text.strip()
    ]
