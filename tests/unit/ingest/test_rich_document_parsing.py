from __future__ import annotations

from decimal import Decimal
from pathlib import Path
from typing import ClassVar

import docx

from app.ingest.parsers import text


def test_docx_table_keeps_table_and_cell_reference(tmp_path: Path) -> None:
    path = tmp_path / "report.docx"
    document = docx.Document()
    document.add_paragraph("经营数据")
    table = document.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "指标"
    table.rows[0].cells[1].text = "本期"
    table.rows[0].cells[2].text = "同比"
    table.rows[1].cells[0].text = "营业收入"
    table.rows[1].cells[1].text = "120亿元"
    table.rows[1].cells[2].text = "增长20%"
    document.save(path)

    parsed = text.parse_docx(path)

    row = next(
        item
        for item in parsed.segments
        if item.content_kind == "table_row" and "120" in item.content
    )
    assert row.table_index == 1
    assert row.row_index == 2
    assert row.cell_range == "A2:C2"
    assert row.extraction_method == "native"


def test_scanned_pdf_uses_ocr_and_keeps_page_confidence(tmp_path: Path, monkeypatch) -> None:
    path = tmp_path / "scan.pdf"
    path.write_bytes(b"placeholder")

    class Page:
        def extract_text(self) -> str:
            return ""

    class Reader:
        pages: ClassVar[list[Page]] = [Page()]

    class PlumberPage:
        def extract_tables(self):
            return []

    class Plumber:
        pages: ClassVar[list[PlumberPage]] = [PlumberPage()]

        def __enter__(self):
            return self

        def __exit__(self, *args):
            return None

    monkeypatch.setattr("pypdf.PdfReader", lambda value: Reader())
    monkeypatch.setattr("pdfplumber.open", lambda value: Plumber())
    monkeypatch.setattr(
        text,
        "_ocr_pdf_page",
        lambda value, page: [("扫描页营业收入同比增长20%", Decimal("0.91"))],
    )

    parsed = text.parse_pdf(path)

    assert parsed.segments[0].page == 1
    assert parsed.segments[0].extraction_method == "ocr"
    assert parsed.segments[0].confidence == Decimal("0.91")
